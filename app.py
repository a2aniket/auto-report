import streamlit as st
import docx
import os
import json
from datetime import datetime, date

def replace_text_in_paragraph(paragraph, replacements):
    if hasattr(paragraph, 'runs'):
        for key, value in replacements.items():
            if key in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        inline[i].text = inline[i].text.replace(key, value)

def replace_text_in_document(doc, replacements):
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

def save_data_to_json(data, filename):
    data_folder = "data"
    if not os.path.exists(data_folder):
        os.makedirs(data_folder)
    filepath = os.path.join(data_folder, filename)
    with open(filepath, "w") as file:
        json.dump(data, file)

def load_data_from_json(filename):
    data_folder = "data"
    filepath = os.path.join(data_folder, filename)
    if os.path.exists(filepath):
        with open(filepath, "r") as file:
            return json.load(file)
    return None

def delete_saved_file(filename):
    data_folder = "data"
    filepath = os.path.join(data_folder, filename)
    if os.path.exists(filepath):
        os.remove(filepath)

def convert_to_date(date_str):
    """Converts a date string to a date object if not None."""
    if date_str:
        return datetime.strptime(date_str, "%d/%m/%Y").date()
    return None

def main():
    st.title("Word Template Modifier")

    # Create folders if they don't exist
    report_folder = "reports"
    download_folder = "downloads"
    data_folder = "data"
    for folder in [report_folder, download_folder, data_folder]:
        if not os.path.exists(folder):
            os.makedirs(folder)

    template_files = [f for f in os.listdir(report_folder) if f.endswith(".docx")]
    saved_files = [f for f in os.listdir(data_folder) if f.endswith(".json")]

    if not template_files:
        st.error("No Word templates found in the 'reports' folder.")
        return

    # Sidebar for saved files
    st.sidebar.title("Saved Reports")
    selected_saved_file = st.sidebar.selectbox("Select a saved report:", [""] + saved_files)

    # Load saved data if a saved file is selected
    saved_data = {}
    if selected_saved_file:
        saved_data = load_data_from_json(selected_saved_file)

    # Template selection
    selected_template = st.selectbox("Select a template:", template_files)

    # Input Fields with saved data pre-fill
    member_code = st.text_input("Member Code", value=saved_data.get("member_code", ""))
    member_name = st.text_input("Member Name", value=saved_data.get("member_name", ""))
    category = st.text_input("Category", value=saved_data.get("category", ""))
    software_name = st.text_input("Software Name", value=saved_data.get("software_name", ""))
    segment = st.text_input("Segment", value=saved_data.get("segment", ""))
    version = st.text_input("Version", value=saved_data.get("version", ""))
    strategy_name = st.text_input("Strategy Name", value=saved_data.get("strategy_name", ""))
    browser_exe_lite = st.text_input("Browser/Exe/Lite", value=saved_data.get("browser_exe_lite", ""))
    vendor_name = st.text_input("Vendor Name", value=saved_data.get("vendor_name", ""))

    if selected_template == "NSE_Report_IBT.docx":
        location_ind = st.text_input("Location India", value=saved_data.get("location_ind", ""))
        net_dig = st.text_input("Network Diagram", value=saved_data.get("net_dig", ""))
        ogv = st.text_input("Order Gateway Version", value=saved_data.get("ogv", ""))
        fev = st.text_input("Frontend Version", value=saved_data.get("fev", ""))
        trader_id = st.text_input("Trader ID", value=saved_data.get("trader_id", ""))
        segment_name = st.text_input("Market Segment", value=saved_data.get("segment_name", ""))
        ip_address = st.text_input("IP Address", value=saved_data.get("ip_address", ""))
        leas_id = st.text_input("Leased Line ID", value=saved_data.get("leas_id", ""))
        firewall_name = st.text_input("Firewall Name", value=saved_data.get("firewall_name", ""))
    else:
        location_ind = ""
        net_dig = ""
        ogv = ""
        fev = ""
        trader_id = ""
        segment_name = ""
        ip_address = ""
        leas_id = ""
        firewall_name = ""

    # Fix for date input issue
    uat_date = st.date_input("UAT Date", value=convert_to_date(saved_data.get("UAT_date", "")) or date.today())
    test_date = st.date_input("Test Date", value=convert_to_date(saved_data.get("test_date", "")) or date.today())
    mock_date = st.date_input("Mock Date", value=convert_to_date(saved_data.get("mock_date", "")) or date.today())

    sor = st.text_input("SOR", value=saved_data.get("SOR", ""))
    rms_name = st.text_input("RMS Name", value=saved_data.get("RMS_name", ""))

    # Define placeholders in template
    replacements = {
        "member_code": member_code,
        "member_name": member_name,
        "category": category,
        "software_name": software_name,
        "segment": segment,
        "version": version,
        "strategy_name": strategy_name,
        "browser_exe_lite": browser_exe_lite,
        "UAT_date": uat_date.strftime("%d/%m/%Y") if uat_date else "",
        "test_date": test_date.strftime("%d/%m/%Y") if test_date else "",
        "mock_date": mock_date.strftime("%d/%m/%Y") if mock_date else "",
        "SOR": sor,
        "RMS_name": rms_name,
        "vendor_name": vendor_name,
        "location_ind": location_ind,
        "net_dig": net_dig,
        "ogv": ogv,
        "fev": fev,
        "trader_id": trader_id,
        "segment_name": segment_name,
        "ip_address": ip_address,
        "leas_id": leas_id,
        "firewall_name": firewall_name
    }

    # Check mandatory fields
    if not member_code or not member_name:
        st.error("Member Code and Member Name are required.")
        return

    # Check for missing fields
    missing_fields = [key for key, value in replacements.items() if not value]

    # Replace values in template if button clicked
    if st.button("Replace Values in Template"):
        if missing_fields:
            # Save data only if fields are missing
            filename = f"{member_code}_{selected_template.split('.')[0]}.json"
            save_data_to_json(replacements, filename)
            st.sidebar.warning("Missing Fields: " + ", ".join(missing_fields))
            st.sidebar.info("Saved current input for later.")
        else:
            # Delete saved report if all fields are filled
            if selected_saved_file and saved_data.get("member_code") == member_code:
                delete_saved_file(selected_saved_file)
                st.sidebar.success(f"Deleted saved report: {selected_saved_file}")

        # Proceed with replacing values
        try:
            doc = docx.Document(os.path.join(report_folder, selected_template))
            replace_text_in_document(doc, replacements)

            # Save modified document in 'downloads' folder
            output_filename = os.path.join(download_folder, f"transfer_{member_name}.docx")
            doc.save(output_filename)
            st.success("Values replaced successfully in the template.")
        except Exception as e:
            st.error(f"Error: {e}")

    # Download button for the modified document
    output_filepath = os.path.join(download_folder, f"transfer_{member_name}.docx")
    if os.path.exists(output_filepath):
        with open(output_filepath, "rb") as file:
            st.download_button(
                label="Download Modified Document",
                data=file,
                file_name=f"transfer_{member_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

if __name__ == "__main__":
    main()
